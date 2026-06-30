from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from textwrap import wrap
from typing import Literal

ExportFormat = Literal["docx", "pdf"]


@dataclass(frozen=True)
class AnswerExport:
    data: bytes
    filename: str
    media_type: str


def requested_export_format(prompt: str) -> ExportFormat | None:
    matches: list[tuple[int, ExportFormat]] = []
    docx_match = re.search(r"\b(docx|word document|microsoft word)\b|\.docx\b", prompt, re.IGNORECASE)
    pdf_match = re.search(r"\bpdf\b|\.pdf\b", prompt, re.IGNORECASE)
    if docx_match:
        matches.append((docx_match.start(), "docx"))
    if pdf_match:
        matches.append((pdf_match.start(), "pdf"))
    if not matches:
        return None
    return sorted(matches, key=lambda item: item[0])[0][1]


def build_answer_export(
    *,
    export_format: ExportFormat,
    title: str,
    answer: str,
    citations: list[dict],
) -> AnswerExport:
    filename = f"{_safe_filename(title)}.{export_format}"
    if export_format == "docx":
        return AnswerExport(
            data=_build_docx(title, answer, citations),
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    return AnswerExport(data=_build_pdf(title, answer, citations), filename=filename, media_type="application/pdf")


def _build_docx(title: str, answer: str, citations: list[dict]) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading(title, level=1)
    for block in _paragraphs(answer):
        document.add_paragraph(block)
    if citations:
        document.add_heading("Sources", level=2)
        for citation in citations:
            source = _citation_label(citation)
            preview = str(citation.get("preview") or "").strip()
            paragraph = document.add_paragraph()
            paragraph.add_run(source).bold = True
            if preview:
                paragraph.add_run(f"\n{preview}")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_pdf(title: str, answer: str, citations: list[dict]) -> bytes:
    import fitz

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    y = 48.0
    y = _write_pdf_block(page, title, x=48, y=y, size=16, bold=True)
    y += 12
    for block in _paragraphs(answer):
        y = _write_pdf_block(page, block, x=48, y=y, size=10.5)
        y += 8
        if y > 760:
            page = doc.new_page(width=595, height=842)
            y = 48.0
    if citations:
        if y > 700:
            page = doc.new_page(width=595, height=842)
            y = 48.0
        y = _write_pdf_block(page, "Sources", x=48, y=y + 8, size=13, bold=True)
        y += 8
        for citation in citations:
            text = _citation_label(citation)
            preview = str(citation.get("preview") or "").strip()
            if preview:
                text = f"{text}\n{preview}"
            y = _write_pdf_block(page, text, x=48, y=y, size=9)
            y += 8
            if y > 760:
                page = doc.new_page(width=595, height=842)
                y = 48.0
    data = doc.tobytes()
    doc.close()
    return data


def _write_pdf_block(page, text: str, *, x: float, y: float, size: float, bold: bool = False) -> float:
    line_height = size * 1.35
    lines: list[str] = []
    for raw_line in text.splitlines() or [""]:
        wrapped = wrap(raw_line, width=max(30, int(86 * 10 / size))) or [""]
        lines.extend(wrapped)
    for line in lines:
        page.insert_text((x, y), line, fontsize=size, fontname="helv", color=(0.09, 0.09, 0.11))
        y += line_height
    return y


def _paragraphs(text: str) -> list[str]:
    return [part.strip() for part in re.split(r"\n\s*\n", text.strip()) if part.strip()]


def _citation_label(citation: dict) -> str:
    source_id = citation.get("id") or "?"
    name = citation.get("document_name") or "Source"
    source_type = citation.get("source_type") or "Indexed KB"
    url = citation.get("source_url")
    return f"[{source_id}] {name} ({source_type}){f' - {url}' if url else ''}"


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", value.strip()).strip("-._")
    return (cleaned or "answer")[:80]
